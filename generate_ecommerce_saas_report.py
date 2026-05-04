from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_FILE = "ecommerce-saas-architecture.docx"


PROJECT_NAME = "E-Commerce SaaS Platforma"
SUBTITLE = "Gamybinei aplinkai paruoštos e. komercijos sistemos techninis ir verslo planas"
AUTHOR = "AI Generated"
REPORT_DATE = "2026-05-04"


THEME = {
    "navy": "163A5F",
    "blue": "1F4E79",
    "slate": "334155",
    "muted": "64748B",
    "gold": "B0892F",
    "light_blue": "EAF2F8",
    "light_gray": "F3F6F8",
    "border": "C9D3DF",
    "white": "FFFFFF",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Psl. ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, fld_text, fld_end])


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(THEME["slate"])
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for idx, size in [(1, 20), (2, 15), (3, 12.5)]:
        style = styles[f"Heading {idx}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(THEME["navy"] if idx == 1 else THEME["blue"])
        style.paragraph_format.space_before = Pt(18 if idx == 1 else 10)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "E-Commerce SaaS Platforma | Techninis ir verslo planas"
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor.from_string(THEME["muted"])
    add_page_number(footer.add_paragraph())


def add_title_page(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("ENTERPRISE TECHNICAL BLUEPRINT")
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(THEME["gold"])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(PROJECT_NAME)
    r.font.name = "Aptos Display"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(THEME["navy"])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(SUBTITLE)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(THEME["slate"])

    meta = [
        ("Autorius", AUTHOR),
        ("Data", REPORT_DATE),
        ("Dokumento tipas", "Techninė ir verslo ataskaita"),
        ("Auditorija", "Kūrėjai, vadovai, produkto savininkai, investuotojai"),
        ("Apimtis", "JWT, RBAC, produktai, krepšelis, užsakymai, Stripe, prenumeratos, PDF sąskaitos, refundai, MongoDB Atlas"),
    ]
    add_table(document, ["Laukas", "Reikšmė"], meta, widths=[1.8, 5.7])

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(20)
    note.paragraph_format.space_after = Pt(20)
    run = note.add_run(
        "Dokumentas parengtas kaip aukšto lygio SaaS architektūros planas. "
        "Jame sąmoningai nepateikiamas pilnas programos kodas; vietoje to aprašomi moduliai, "
        "duomenų modeliai, API, saugumo sluoksniai, mokėjimų srautai ir diegimo reikalavimai."
    )
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(THEME["muted"])

    document.add_page_break()


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = str(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, THEME["blue"])
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(THEME["white"])
                run.font.size = Pt(9.5)

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, THEME["light_gray"])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string(THEME["slate"])

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = text
    set_cell_shading(cell, THEME["light_blue"])
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(8.8)
            run.font.color.rgb = RGBColor.from_string(THEME["slate"])
    document.add_paragraph()


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FFF7E6")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("8A5A00")
    r.font.size = Pt(10)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(THEME["slate"])
        run.font.size = Pt(9.5)
    document.add_paragraph()


def add_toc(document):
    document.add_heading("Turinys", level=1)
    rows = [
        ("1", "Vadovybės santrauka"),
        ("2", "Sistemos architektūra"),
        ("3", "Autentifikacijos sistema (JWT)"),
        ("4", "Duomenų bazės dizainas"),
        ("5", "API dizainas"),
        ("6", "Administratoriaus panelė"),
        ("7", "E. komercijos srautas"),
        ("8", "Stripe integracija"),
        ("9", "PDF sąskaitų sistema"),
        ("10", "Grąžinimų ir atšaukimų sistema"),
        ("11", "Saugumo sluoksnis"),
        ("12", "Diegimo strategija"),
        ("13", "Kūrimo kelrodis"),
        ("14", "Sudėtingumo ir laiko įvertinimas"),
        ("15", "Galutinis kontrolinis sąrašas"),
    ]
    add_table(document, ["Nr.", "Skyrius"], rows, widths=[0.7, 6.8])
    document.add_page_break()


def build_report():
    document = Document()
    configure_document(document)
    add_title_page(document)
    add_toc(document)

    document.add_heading("1. Vadovybės santrauka", level=1)
    document.add_paragraph(
        "Šis dokumentas apibrėžia, kaip sukurti pilną, gamybinei aplinkai paruoštą e. komercijos SaaS platformą, "
        "kuri sujungia klientų apsipirkimą, administratoriaus valdymą, saugią JWT autentifikaciją, Stripe mokėjimus, "
        "prenumeratas, PDF sąskaitas, refundų valdymą ir MongoDB Atlas duomenų sluoksnį."
    )
    document.add_paragraph(
        "Tikslas yra ne tik sukurti internetinę parduotuvę, bet ir parengti sistemą, kurią galima eksploatuoti kaip "
        "komercinį SaaS produktą: su aiškiomis rolėmis, mokėjimų įvykių auditu, plečiama architektūra, observability, "
        "saugumo kontrole ir skaidriu diegimo procesu."
    )
    add_table(
        document,
        ["Aspektas", "Vertė verslui"],
        [
            ("Klientų savitarna", "Registracija, prisijungimas, krepšelis, užsakymai, sąskaitos ir prenumeratos mažina rankinį aptarnavimą."),
            ("Pasikartojančios pajamos", "Stripe prenumeratos leidžia kurti mėnesinius ir metinius planus, trial periodus ir SaaS pajamų modelį."),
            ("Operacinis valdymas", "Admin panelė leidžia valdyti produktus, užsakymus, mokėjimus, refundus ir audito įrašus vienoje vietoje."),
            ("Investicinis patrauklumas", "Aiški architektūra, mokėjimų kontrolė ir diegimo planas mažina techninę riziką ir leidžia prognozuoti plėtrą."),
        ],
        widths=[2.2, 5.3],
    )
    document.add_heading("Tiksliniai naudotojai", level=2)
    add_bullets(
        document,
        [
            "Klientai, kurie perka vienkartinius produktus arba užsisako prenumeratas.",
            "Administratoriai, valdantys produktų katalogą, atsargas, kainas, užsakymus ir refundus.",
            "Verslo vadovai, stebintys pardavimus, prenumeratų augimą ir finansinę būseną.",
            "Kūrėjų komanda, kuri palaiko API, duomenų bazę, integracijas ir diegimą.",
        ],
    )

    document.add_heading("2. Sistemos architektūra", level=1)
    document.add_paragraph(
        "Rekomenduojama architektūra yra modulinė: React/Vite frontend, Node.js/Express backend, MongoDB Atlas duomenų bazė, "
        "Stripe mokėjimų sluoksnis ir atskiras failų saugyklos sluoksnis produktų nuotraukoms bei PDF sąskaitoms. "
        "Sistema gali būti pradėta kaip modulinis monolitas, tačiau moduliai turi būti atskirti taip, kad ateityje mokėjimus, "
        "sąskaitas ar katalogą būtų galima iškelti į savarankiškas paslaugas."
    )
    add_table(
        document,
        ["Sluoksnis", "Technologija", "Atsakomybė"],
        [
            ("Frontend", "React + Vite", "Kliento parduotuvė, admin panelė, formos, krepšelis, checkout inicijavimas."),
            ("Backend", "Node.js + Express", "REST API, JWT, RBAC, verslo taisyklės, Stripe sesijos, webhookai, PDF generavimas."),
            ("Database", "MongoDB Atlas + Mongoose", "Vartotojai, produktai, užsakymai, mokėjimai, prenumeratos, sąskaitos."),
            ("Payment", "Stripe Checkout + Billing", "Vienkartiniai mokėjimai, prenumeratos, refundai, mokėjimo patvirtinimai."),
            ("Storage", "S3/Cloudinary arba analogas", "Produktų nuotraukos ir privačios PDF sąskaitos."),
            ("Deployment", "Vercel + Railway/Render", "Atskiras frontend ir backend diegimas, aplinkų valdymas, SSL."),
        ],
        widths=[1.6, 2.2, 3.7],
    )
    document.add_heading("Architektūros srauto diagrama", level=2)
    add_code_block(
        document,
        """React/Vite Frontend
  -> REST API request
      -> Express Backend
          -> JWT middleware
          -> RBAC policy layer
          -> Domain modules
              -> Products / Cart / Orders
              -> Payments / Subscriptions
              -> Invoices / Audit Logs
          -> MongoDB Atlas
          -> Stripe API
          -> File Storage

Stripe
  -> Webhook endpoint
      -> Signature validation
      -> Idempotent event processing
      -> Payment/order/subscription update
      -> Invoice generation""",
    )
    add_callout(
        document,
        "Architektūrinė rekomendacija",
        "Kainos, atsargos, mokesčiai ir mokėjimų patvirtinimai turi būti valdomi backend pusėje. Frontend niekada neturi būti vienintelis tiesos šaltinis finansinėms reikšmėms.",
    )

    document.add_heading("3. Autentifikacijos sistema (JWT)", level=1)
    document.add_paragraph(
        "Autentifikacijos sistema turi palaikyti klientų ir administratorių roles. Access token naudojamas trumpalaikei API prieigai, "
        "refresh token naudojamas sesijai pratęsti. Refresh token turi būti saugomas httpOnly, Secure cookie formatu, o duomenų bazėje laikomas tik jo hash."
    )
    document.add_heading("Login/register srautas", level=2)
    add_numbered(
        document,
        [
            "Vartotojas pateikia el. paštą ir slaptažodį registracijos arba prisijungimo formoje.",
            "Backend validuoja duomenis, patikrina slaptažodį ir vartotojo statusą.",
            "Sėkmės atveju sugeneruojamas trumpalaikis access token ir ilgesnio galiojimo refresh token.",
            "Refresh token hash įrašomas į vartotojo sesiją duomenų bazėje.",
            "Klientas gauna access token ir httpOnly cookie su refresh token.",
            "Kiekvienas saugomas API endpointas tikrina access token ir rolės leidimus.",
        ],
    )
    document.add_heading("Rolės ir leidimai", level=2)
    add_table(
        document,
        ["Rolė", "Leidimai"],
        [
            ("customer", "Registruotis, prisijungti, valdyti krepšelį, kurti užsakymus, matyti savo sąskaitas ir prenumeratas."),
            ("admin", "Valdyti produktus, užsakymus, mokėjimus, refundus, prenumeratas, sąskaitas ir audito įrašus."),
            ("super_admin", "Pasirinktina rolė: valdyti kitus administratorius, jautrius nustatymus ir platformos konfigūraciją."),
        ],
        widths=[1.7, 5.8],
    )
    document.add_heading("Saugumo praktikos", level=2)
    add_bullets(
        document,
        [
            "Naudoti bcrypt arba argon2 slaptažodžių hashinimui.",
            "Access token galiojimas: 5-15 minučių; refresh token galiojimas: pagal verslo riziką, pavyzdžiui 7-30 dienų.",
            "Refresh token rotuoti kiekvieno atnaujinimo metu ir aptikti pakartotinį panaudojimą.",
            "Admin prisijungimui naudoti griežtesnį rate limiting ir rekomenduoti MFA.",
            "Frontend route guard naudojamas UX tikslais, tačiau tikras leidimų tikrinimas vyksta backend pusėje.",
        ],
    )

    document.add_heading("4. Duomenų bazės dizainas", level=1)
    document.add_paragraph(
        "Duomenų bazės schema turi atspindėti verslo domeną. Užsakymų dokumentuose būtina išsaugoti kainas ir produktų pavadinimus pirkimo momentu, "
        "nes produktų katalogas ateityje keisis. Mokėjimų ir prenumeratų kolekcijos turi saugoti Stripe identifikatorius, kad webhookai galėtų tiksliai susieti įvykius."
    )
    add_table(
        document,
        ["Kolekcija", "Pagrindiniai laukai", "Pastabos"],
        [
            ("Users", "email, passwordHash, role, status, name, addresses, stripeCustomerId, refreshTokenHash", "Unikalus email, rolės customer/admin, sesijų valdymas."),
            ("Products", "sku, slug, name, description, images, price, currency, stock, category, status", "Unikalūs sku ir slug, indeksai paieškai bei filtrams."),
            ("Orders", "orderNumber, userId, items, subtotal, tax, total, status, paymentId, invoiceId", "Užsakymo momentu saugoti kainas, mokesčius ir produktų snapshot."),
            ("Payments", "orderId, stripeSessionId, stripePaymentIntentId, amount, currency, status, refundStatus", "Stripe įvykių suderinimas, refund istorija."),
            ("Subscriptions", "userId, planId, stripeSubscriptionId, status, currentPeriodStart, currentPeriodEnd", "Mėnesiniai/metiniai planai, cancelAtPeriodEnd."),
            ("Invoices", "invoiceNumber, orderId, userId, pdfUrl, subtotal, tax, total, issuedAt", "Privatus PDF, atsisiuntimas tik savininkui arba adminui."),
            ("AuditLogs", "actorId, action, targetType, targetId, before, after, ip, userAgent", "Admin veiksmų ir finansinių operacijų atsekamumas."),
        ],
        widths=[1.5, 4.1, 1.9],
    )

    document.add_heading("5. API dizainas", level=1)
    document.add_paragraph(
        "REST API turi būti aiškiai padalintas į auth, shop, admin, checkout, subscription ir webhook maršrutus. Visi admin maršrutai privalo naudoti admin RBAC guard."
    )
    add_table(
        document,
        ["Metodas", "Endpoint", "Prieiga", "Paskirtis"],
        [
            ("POST", "/api/auth/register", "Public", "Kliento registracija."),
            ("POST", "/api/auth/login", "Public", "Prisijungimas ir tokenų išdavimas."),
            ("POST", "/api/auth/refresh", "Refresh cookie", "Access token atnaujinimas ir refresh token rotacija."),
            ("POST", "/api/auth/logout", "Authenticated", "Sesijos nutraukimas."),
            ("GET", "/api/products", "Public", "Produktų sąrašas, filtrai, paieška."),
            ("GET", "/api/products/:slug", "Public", "Produkto detalės."),
            ("POST", "/api/admin/products", "Admin", "Sukurti produktą."),
            ("PATCH", "/api/admin/products/:id", "Admin", "Redaguoti produktą."),
            ("DELETE", "/api/admin/products/:id", "Admin", "Archyvuoti arba ištrinti produktą."),
            ("POST", "/api/cart/items", "Customer", "Pridėti prekę į krepšelį."),
            ("PATCH", "/api/cart/items/:id", "Customer", "Atnaujinti kiekį."),
            ("DELETE", "/api/cart/items/:id", "Customer", "Pašalinti prekę."),
            ("POST", "/api/orders", "Customer", "Sukurti užsakymą iš krepšelio."),
            ("POST", "/api/checkout/session", "Customer", "Sukurti Stripe Checkout Session vienkartiniam mokėjimui."),
            ("POST", "/api/subscriptions/session", "Customer", "Sukurti Stripe Checkout Session prenumeratai."),
            ("POST", "/api/stripe/webhook", "Stripe", "Apdoroti Stripe mokėjimų ir prenumeratų įvykius."),
            ("POST", "/api/admin/payments/:id/refund", "Admin", "Inicijuoti refundą."),
            ("POST", "/api/admin/orders/:id/cancel", "Admin", "Atšaukti užsakymą."),
            ("GET", "/api/invoices/:id/download", "Owner/Admin", "Atsisiųsti PDF sąskaitą."),
        ],
        widths=[0.8, 2.5, 1.3, 2.9],
    )

    document.add_heading("6. Administratoriaus panelė", level=1)
    document.add_paragraph(
        "Admin panelė turi būti operacinis valdymo centras. Jos tikslas - greitai priimti sprendimus, saugiai atlikti finansinius veiksmus ir išlaikyti aiškų audito pėdsaką."
    )
    add_table(
        document,
        ["Modulis", "Funkcijos"],
        [
            ("Dashboard", "Pardavimų suvestinė, aktyvūs užsakymai, mokėjimo klaidos, žemi likučiai, prenumeratų būsenos."),
            ("Product CRUD", "Produkto sukūrimas, redagavimas, archyvavimas, nuotraukų įkėlimas, kainos, atsargos, kategorijos."),
            ("Order Management", "Užsakymo statusų keitimas, kliento informacija, sąskaitos peržiūra, užsakymo atšaukimas."),
            ("Payment/Refund Management", "Mokėjimo statusai, Stripe identifikatoriai, pilnas arba dalinis refund, refund priežastis."),
            ("Subscription Management", "Aktyvios prenumeratos, planai, statusai, atšaukimas, periodų datos."),
            ("Audit Logs", "Admin veiksmų istorija: kas, kada, ką pakeitė ir iš kurio IP."),
        ],
        widths=[2.0, 5.5],
    )

    document.add_heading("7. E. komercijos srautas", level=1)
    document.add_paragraph(
        "Pirkimo srautas turi būti aiškus klientui ir techniškai deterministinis backend pusėje. Krepšelio turinys gali keistis frontend pusėje, tačiau galutinės kainos ir atsargos tikrinamos serveryje."
    )
    add_code_block(
        document,
        """Cart
  -> Validate stock and prices on backend
  -> Create order with status pending_payment
  -> Create Stripe Checkout Session
  -> Redirect customer to Stripe
  -> Receive Stripe webhook
  -> Mark payment as paid
  -> Mark order as paid/processing
  -> Generate PDF invoice
  -> Customer downloads invoice from account page""",
    )
    add_numbered(
        document,
        [
            "Klientas prideda produktą į krepšelį.",
            "Backend patikrina produkto egzistavimą ir leidžiamą kiekį.",
            "Klientas inicijuoja checkout.",
            "Backend sukuria užsakymą pending_payment būsenoje.",
            "Backend sukuria Stripe Checkout Session pagal serverio kainas.",
            "Stripe webhook patvirtina mokėjimą.",
            "Sistema atnaujina užsakymą, mokėjimą ir sugeneruoja PDF sąskaitą.",
        ],
    )

    document.add_heading("8. Stripe integracija", level=1)
    document.add_paragraph(
        "Stripe integracija turi būti paremta Checkout Session ir webhook patvirtinimu. Success URL yra tik vartotojo nukreipimo vieta, bet ne galutinis mokėjimo įrodymas."
    )
    document.add_heading("Vienkartiniai mokėjimai", level=2)
    add_bullets(
        document,
        [
            "Backend sukuria Checkout Session su mode=payment.",
            "Line items formuojami iš serverio duomenų bazės kainų.",
            "Payment status MongoDB atnaujinamas tik po webhook patvirtinimo.",
            "Webhook eventId saugomas, kad tas pats įvykis nebūtų apdorotas du kartus.",
        ],
    )
    document.add_heading("Prenumeratos", level=2)
    add_bullets(
        document,
        [
            "Mėnesiniai ir metiniai planai saugomi kaip platformos planai ir susiejami su Stripe Price ID.",
            "Checkout Session kuriama su mode=subscription.",
            "customer.subscription.created, customer.subscription.updated, customer.subscription.deleted ir invoice.paid webhookai sinchronizuoja būseną.",
            "Prieiga prie prenumeruojamų funkcijų nustatoma pagal lokalią subscription.status reikšmę.",
        ],
    )
    document.add_heading("Webhook apdorojimas", level=2)
    add_code_block(
        document,
        """Webhook request
  -> Read raw body
  -> Verify Stripe signature
  -> Check if eventId already processed
  -> Route by event type
      checkout.session.completed
      payment_intent.succeeded
      invoice.paid
      invoice.payment_failed
      customer.subscription.updated
      customer.subscription.deleted
      charge.refunded
  -> Update MongoDB
  -> Return 2xx response""",
    )

    document.add_heading("9. PDF sąskaitų sistema", level=1)
    document.add_paragraph(
        "PDF sąskaita generuojama po sėkmingo mokėjimo patvirtinimo. Ji turi būti sukurta iš užsakymo snapshot duomenų, kad vėlesni produktų ar kainų pakeitimai nepakeistų istorinės sąskaitos."
    )
    add_table(
        document,
        ["Duomenys sąskaitoje", "Aprašymas"],
        [
            ("Pardavėjas", "Įmonės pavadinimas, adresas, įmonės kodas, PVM kodas, kontaktai."),
            ("Pirkėjas", "Vardas, pavardė arba įmonė, el. paštas, adresas, PVM kodas jei taikoma."),
            ("Užsakymas", "Užsakymo numeris, data, mokėjimo statusas, Stripe PaymentIntent ID."),
            ("Prekės", "SKU, pavadinimas, kiekis, vieneto kaina, nuolaida, mokestis, eilutės suma."),
            ("Mokesčiai", "PVM/tax tarifas, mokesčio suma, bendra suma."),
            ("Failas", "Privatus PDF URL arba saugyklos raktas, atsisiunčiamas per autorizuotą endpointą."),
        ],
        widths=[2.2, 5.3],
    )
    add_numbered(
        document,
        [
            "Webhook patvirtina sėkmingą mokėjimą.",
            "Sistema patikrina, ar užsakymas dar neturi sąskaitos.",
            "Sugeneruojamas unikalus invoiceNumber.",
            "PDF sukuriamas pagal HTML arba PDFKit šabloną.",
            "PDF įkeliamas į privačią saugyklą.",
            "Invoices kolekcijoje sukuriamas įrašas ir susiejamas su užsakymu.",
        ],
    )

    document.add_heading("10. Grąžinimų ir atšaukimų sistema", level=1)
    document.add_paragraph(
        "Refund ir cancel veiksmai turi būti atliekami tik admin panelėje, su aiškia priežastimi, audit log įrašu ir Stripe atsakymo saugojimu. "
        "Dalinis refund turi atnaujinti ir mokėjimo, ir užsakymo statusą."
    )
    add_code_block(
        document,
        """Admin selects payment/order
  -> Enters refund amount and reason
  -> Backend validates admin role and allowed state
  -> Create audit log: refund_requested
  -> Call Stripe Refund API with idempotency key
  -> Update payment refundStatus
  -> Update order status
  -> Create audit log: refund_completed
  -> Notify customer if business policy requires it""",
    )
    add_table(
        document,
        ["Veiksmas", "Rezultatas sistemoje"],
        [
            ("Cancel unpaid order", "Order status=canceled, Payment status=canceled arba abandoned."),
            ("Full refund", "Payment refundStatus=refunded, Order status=refunded, audit log sukurtas."),
            ("Partial refund", "Payment refundStatus=partially_refunded, Order status=partially_refunded."),
            ("Subscription cancel", "cancelAtPeriodEnd=true arba immediate cancel pagal pasirinktą politiką."),
        ],
        widths=[2.2, 5.3],
    )

    document.add_heading("11. Saugumo sluoksnis", level=1)
    add_table(
        document,
        ["Sritis", "Reikalavimas"],
        [
            ("JWT storage", "Refresh token saugoti httpOnly Secure cookie, DB laikyti tik hash. Access token trumpalaikis."),
            ("Password hashing", "bcrypt arba argon2 su tinkamu cost parametru. Slaptažodžių neatvaizduoti loguose."),
            ("API protection", "JWT middleware, RBAC guard, request schema validacija, server-side kainų perskaičiavimas."),
            ("Rate limiting", "Atskiri limitai login, register, refresh, checkout, admin ir webhook endpointams."),
            ("CSRF/CORS", "Aiški CORS politika, CSRF apsauga cookie pagrindu veikiančiam refresh srautui."),
            ("Stripe security", "Webhook signature verification, raw body, event idempotency, secret keys tik serverio aplinkoje."),
            ("Audit", "Admin finansiniai ir produktų pakeitimai saugomi auditLogs kolekcijoje."),
            ("Secrets", "Naudoti hostingo secrets manager, niekada necommitinti .env su realiais raktais."),
        ],
        widths=[2.0, 5.5],
    )

    document.add_heading("12. Diegimo strategija", level=1)
    document.add_paragraph(
        "Diegimas turi būti automatizuotas ir atskirtas pagal aplinkas: development, staging ir production. Staging turi naudoti Stripe test mode ir atskirą MongoDB duomenų bazę."
    )
    add_table(
        document,
        ["Komponentas", "Platforma", "Pastabos"],
        [
            ("Frontend", "Vercel", "React/Vite statinis build, env vars tik public reikšmėms, SSL ir CDN."),
            ("Backend", "Railway arba Render", "Node.js/Express serveris, private env vars, health check, logs."),
            ("Database", "MongoDB Atlas", "Atskiri vartotojai, IP allowlist, backups, monitoringas."),
            ("Payments", "Stripe", "Test ir live raktai, atskiri webhook endpointai kiekvienai aplinkai."),
            ("Storage", "S3/Cloudinary", "Produktų vaizdai vieši, PDF sąskaitos privačios."),
        ],
        widths=[1.8, 2.0, 3.7],
    )
    document.add_heading("Svarbiausi aplinkos kintamieji", level=2)
    add_code_block(
        document,
        """NODE_ENV=production
FRONTEND_URL=https://app.example.com
MONGODB_URI=mongodb+srv://...
JWT_ACCESS_SECRET=...
JWT_REFRESH_SECRET=...
STRIPE_SECRET_KEY=sk_live_...
STRIPE_WEBHOOK_SECRET=whsec_...
STORAGE_BUCKET=...
EMAIL_API_KEY=...""",
    )

    document.add_heading("13. Kūrimo kelrodis", level=1)
    add_table(
        document,
        ["Fazė", "Tikslas", "Rezultatas"],
        [
            ("MVP 1", "Repo, React/Vite, Express, MongoDB, bazinis dizainas", "Veikianti projekto struktūra ir CI pagrindas."),
            ("MVP 2", "JWT auth, customer/admin rolės, protected routes", "Saugus prisijungimas ir RBAC."),
            ("MVP 3", "Product CRUD, katalogas, admin produktų valdymas", "Administruojamas produktų katalogas."),
            ("MVP 4", "Cart, checkout, orders", "Pilnas pirkimo srautas iki mokėjimo inicijavimo."),
            ("MVP 5", "Stripe one-time payments ir webhooks", "Patikimas mokėjimo patvirtinimas ir order statusai."),
            ("Full 1", "Subscriptions monthly/yearly", "SaaS prenumeratų modelis."),
            ("Full 2", "PDF invoices", "Sąskaitų generavimas ir atsisiuntimas."),
            ("Full 3", "Refund/cancel, audit logs", "Finansinių veiksmų valdymas admin panelėje."),
            ("Launch", "Security hardening, tests, staging, production", "Gamybinis paleidimas."),
        ],
        widths=[1.4, 3.0, 3.1],
    )

    document.add_heading("14. Sudėtingumo ir laiko įvertinimas", level=1)
    document.add_paragraph(
        "Bendras projekto sudėtingumas yra aukštas, nes sistema apima finansines operacijas, prenumeratų gyvavimo ciklą, administratoriaus veiksmų auditą ir production diegimą."
    )
    add_table(
        document,
        ["Modulis", "Sudėtingumas", "Įvertis"],
        [
            ("JWT + RBAC", "Vidutinis", "1-2 savaitės"),
            ("Product CRUD + Admin UI", "Vidutinis", "1-2 savaitės"),
            ("Cart + Orders", "Vidutinis", "1-2 savaitės"),
            ("Stripe one-time payments", "Vidutinis-aukštas", "1-2 savaitės"),
            ("Subscriptions", "Aukštas", "1-2 savaitės"),
            ("PDF invoices", "Vidutinis", "1 savaitė"),
            ("Refund/cancel + audit", "Aukštas", "1 savaitė"),
            ("Security + deployment", "Aukštas", "1-2 savaitės"),
        ],
        widths=[3.0, 2.0, 2.5],
    )
    document.add_paragraph(
        "Realistiškas pilnos sistemos terminas: 8-12 savaičių patyrusiai komandai, jei dizainas ir verslo taisyklės yra aiškūs. "
        "Papildomai rekomenduojama suplanuoti 1 stabilizavimo savaitę prieš viešą paleidimą."
    )

    document.add_heading("15. Galutinis kontrolinis sąrašas", level=1)
    add_bullets(
        document,
        [
            "Customer ir admin login veikia su access/refresh token logika.",
            "RBAC saugo visus admin ir jautrius API maršrutus.",
            "Product CRUD palaiko kainas, atsargas, nuotraukas ir statusus.",
            "Krepšelis perskaičiuojamas serveryje prieš checkout.",
            "Užsakymas sukuriamas pending_payment būsenoje prieš Stripe nukreipimą.",
            "Stripe webhookai validuoja parašą ir apdoroja įvykius idempotentiškai.",
            "Vienkartiniai mokėjimai ir prenumeratos įrašomos į MongoDB.",
            "PDF sąskaita generuojama po sėkmingo mokėjimo.",
            "Admin gali atlikti refund/cancel su priežastimi ir audito įrašu.",
            "MongoDB Atlas turi backup, indeksus ir atskiras aplinkas.",
            "Rate limiting, validacija, CORS/CSRF ir secrets valdymas įjungti.",
            "Frontend įdiegtas Vercel, backend Railway/Render, production env vars sukonfigūruoti.",
            "Staging aplinkoje ištestuoti Stripe testiniai sėkmės, klaidos, refund ir subscription scenarijai.",
            "Monitoringas, klaidų sekimas ir incidentų eskalavimo tvarka paruošta.",
        ],
    )

    document.core_properties.title = PROJECT_NAME
    document.core_properties.subject = SUBTITLE
    document.core_properties.author = AUTHOR
    document.core_properties.comments = "Enterprise-level Lithuanian technical and business report for an e-commerce SaaS platform."
    document.core_properties.created = datetime(2026, 5, 4, 12, 0, 0)
    document.core_properties.modified = datetime(2026, 5, 4, 12, 0, 0)

    output_path = Path(__file__).with_name(OUTPUT_FILE)
    document.save(output_path)
    print(f"Generated: {output_path.resolve()}")


if __name__ == "__main__":
    build_report()
